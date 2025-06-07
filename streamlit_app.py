import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ec2_sql_sizing import EC2DatabaseSizingCalculator
import os
import time
from dotenv import load_dotenv

# Load environment variables from .env only if needed
if not os.getenv("AWS_ACCESS_KEY_ID") or not os.getenv("AWS_SECRET_ACCESS_KEY"):
    load_dotenv()

# Configure page
st.set_page_config(
    page_title="Enterprise AWS EC2 SQL Sizing", 
    layout="wide",
    page_icon=":bar_chart:"
)

# Add custom CSS for professional styling
st.markdown("""
<style>
    .reportview-container {
        background: #f0f2f6;
    }
    .sidebar .sidebar-content {
        background: #ffffff;
        box-shadow: 0 0 10px rgba(0,0,0,0.1);
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        font-weight: bold;
        border-radius: 4px;
    }
    .stDownloadButton>button {
        background-color: #2196F3;
        color: white;
        font-weight: bold;
        border-radius: 4px;
    }
    .stAlert {
        border-radius: 4px;
    }
    .st-bb {
        background-color: white;
    }
    .metric-box {
        border: 1px solid #e0e0e0;
        border-radius: 5px;
        padding: 15px;
        margin-bottom: 15px;
        background-color: #ffffff;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .metric-title {
        font-weight: bold;
        margin-bottom: 5px;
        color: #2c3e50;
    }
    .credential-status {
        padding: 10px;
        border-radius: 4px;
        margin-top: 10px;
        font-weight: bold;
    }
    .cred-valid {
        background-color: #d4edda;
        color: #155724;
    }
    .cred-invalid {
        background-color: #f8d7da;
        color: #721c24;
    }
</style>
""", unsafe_allow_html=True)

# App header
st.title("AWS EC2 SQL Server Sizing Calculator")
st.markdown("""
This enterprise-grade tool provides EC2 sizing recommendations for SQL Server workloads based on your on-premise infrastructure metrics.
Recommendations include development, QA, staging, and production environments with detailed cost estimates.
""")

# Initialize calculator
calculator = EC2DatabaseSizingCalculator()

# Sidebar Inputs
with st.sidebar:
    st.header("AWS Configuration")
    
    # Region selection
    region = st.selectbox(
        "AWS Region", 
        ["us-east-1", "us-west-1", "us-west-2"],
        index=0,
        help="Select the AWS region for pricing and deployment"
    )
    
    # Credential information
    st.subheader("AWS Credentials")
    st.markdown("""
    Credentials are sourced from:
    - IAM role (when deployed on AWS)
    - Environment variables
    - ~/.aws/credentials file
    """)
    
    # Check credential status with detailed diagnostics
    cred_status, cred_message = calculator.validate_aws_credentials(region)
    status_class = "cred-valid" if cred_status else "cred-invalid"
    status_text = "✅ " + cred_message if cred_status else "⚠️ " + cred_message
    st.markdown(f'<div class="credential-status {status_class}">{status_text}</div>', unsafe_allow_html=True)
    
    if not cred_status:
        # Show credential debug information
        with st.expander("Credential Debug Information", expanded=True):
            st.write("Current environment variables:")
            st.code(f"""
            AWS_ACCESS_KEY_ID: {'*****' if os.environ.get('AWS_ACCESS_KEY_ID') else 'Not set'}
            AWS_SECRET_ACCESS_KEY: {'*****' if os.environ.get('AWS_SECRET_ACCESS_KEY') else 'Not set'}
            AWS_DEFAULT_REGION: {os.environ.get('AWS_DEFAULT_REGION', 'Not set')}
            """)
            
            st.write("AWS configuration files:")
            # Platform-independent path handling
            if os.name == 'nt':  # Windows
                aws_dir = os.path.join(os.environ.get('USERPROFILE', ''), '.aws')
            else:  # Linux/Mac
                aws_dir = os.path.join(os.path.expanduser("~"), '.aws')
                
            aws_files = [
                os.path.join(aws_dir, "credentials"),
                os.path.join(aws_dir, "config")
            ]
            
            for file in aws_files:
                if os.path.exists(file):
                    st.success(f"Found: {file}")
                    try:
                        with open(file, 'r') as f:
                            content = f.read(500)
                        st.code(content)
                    except Exception as e:
                        st.warning(f"Could not read file: {str(e)}")
                else:
                    st.warning(f"Not found: {file}")
    
    st.header("Input Parameters")
    st.markdown("Enter your current on-premise SQL Server metrics:")
    
    # Input sections in expanders
    with st.expander("Compute Resources", expanded=True):
        on_prem_cores = st.number_input("CPU Cores", min_value=1, value=16, 
                                       help="Total number of CPU cores in your on-premise server")
        peak_cpu_percent = st.slider("Peak CPU Utilization (%)", 0, 100, 65, 
                                    help="Highest observed CPU utilization percentage")
        on_prem_ram_gb = st.number_input("RAM (GB)", min_value=1, value=64, 
                                        help="Total physical RAM in the server")
        peak_ram_percent = st.slider("Peak RAM Utilization (%)", 0, 100, 75, 
                                    help="Highest observed RAM utilization percentage")
    
    with st.expander("Storage & I/O", expanded=True):
        storage_current_gb = st.number_input("Current Storage (GB)", min_value=1, value=500, 
                                           help="Current database storage size")
        storage_growth_rate = st.number_input("Annual Growth Rate", min_value=0.0, max_value=1.0, value=0.15, step=0.01, 
                                             format="%.2f", help="Expected annual storage growth (e.g., 0.15 for 15%)")
        peak_iops = st.number_input("Peak IOPS", min_value=1, value=8000, 
                                   help="Highest observed Input/Output Operations Per Second")
        peak_throughput_mbps = st.number_input("Peak Throughput (MB/s)", min_value=1, value=400, 
                                              help="Highest observed data transfer rate")
    
    with st.expander("Deployment Settings", expanded=True):
        years = st.slider("Growth Projection (Years)", 1, 10, 3, 
                         help="Number of years to plan for future growth")
        workload_profile = st.selectbox("Workload Profile", 
                                      ["general", "memory", "compute"],
                                      help="""\n**Workload Type Guidelines**  \n- General: Balanced workloads like mixed OLTP and reporting  \n- Memory: Data warehouses, analytics, in-memory DBs  \n- Compute: OLTP, heavy transaction processing, CPU-bound jobs\n""")
        prefer_amd = st.checkbox("Include AMD Instances (Cost Optimized)", value=True,
                                help="AMD instances are typically 10-20% cheaper than comparable Intel instances")

# Prepare inputs dictionary
inputs = {
    "region": region,
    "on_prem_cores": on_prem_cores,
    "peak_cpu_percent": peak_cpu_percent,
    "on_prem_ram_gb": on_prem_ram_gb,
    "peak_ram_percent": peak_ram_percent,
    "storage_current_gb": storage_current_gb,
    "storage_growth_rate": storage_growth_rate,
    "peak_iops": peak_iops,
    "peak_throughput_mbps": peak_throughput_mbps,
    "years": years,
    "workload_profile": workload_profile,
    "prefer_amd": prefer_amd
}

# Update calculator inputs
calculator.inputs.update(inputs)

# Main app
if st.button("Generate Recommendations", key="generate_btn"):
    start_time = time.time()
    
    with st.spinner("Calculating EC2 sizing recommendations..."):
        try:
            # Fetch current prices from AWS
            calculator.fetch_current_prices()
            
            # Generate recommendations
            results = calculator.generate_all_recommendations()
            
            # Create DataFrame
            df = pd.DataFrame.from_dict(results, orient='index').reset_index()
            df.rename(columns={"index": "Environment"}, inplace=True)
            
            # Display input summary
            st.subheader("Input Summary")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown('<div class="metric-box">'
                            '<div class="metric-title">Compute</div>'
                            f'Cores: {on_prem_cores}<br>'
                            f'Peak CPU: {peak_cpu_percent}%<br>'
                            f'RAM: {on_prem_ram_gb} GB<br>'
                            f'Peak RAM: {peak_ram_percent}%'
                            '</div>', unsafe_allow_html=True)
            with col2:
                st.markdown('<div class="metric-box">'
                            '<div class="metric-title">Storage</div>'
                            f'Current: {storage_current_gb} GB<br>'
                            f'Growth: {storage_growth_rate*100:.1f}%<br>'
                            f'Projection: {years} years<br>'
                            f'Peak IOPS: {peak_iops:,}'
                            '</div>', unsafe_allow_html=True)
            with col3:
                st.markdown('<div class="metric-box">'
                            '<div class="metric-title">Configuration</div>'
                            f'Region: {region}<br>'
                            f'Workload: {workload_profile.title()}<br>'
                            f'AMD Instances: {"Yes" if prefer_amd else "No"}<br>'
                            '</div>', unsafe_allow_html=True)

            # Display results
            st.subheader("Sizing Recommendations")
            st.success(f"✅ EC2 Sizing Recommendations Generated for {region}")
            
            # Format costs
            formatted_df = df.copy()
            cost_columns = ["instance_cost", "ebs_cost", "total_cost"]
            for col in cost_columns:
                formatted_df[col] = formatted_df[col].apply(lambda x: f"${x:,.2f}")
            
            # Display table
            st.dataframe(
                formatted_df[[
                    "Environment", "instance_type", "vCPUs", "RAM_GB", 
                    "storage_GB", "ebs_type", "total_cost"
                ]],
                use_container_width=True
            )
            
            # Show detailed view
            with st.expander("Detailed View"):
                st.dataframe(
                    formatted_df[[
                        "Environment", "instance_type", "vCPUs", "RAM_GB", 
                        "storage_GB", "ebs_type", "iops_required", 
                        "throughput_required", "family", "processor", 
                        "instance_cost", "ebs_cost", "total_cost"
                    ]],
                    use_container_width=True
                )
            
            # Execution time
            exec_time = time.time() - start_time
            st.caption(f"Execution time: {exec_time:.2f} seconds")
            
            # Export options
            st.subheader("Export Results")
            export_col1, export_col2 = st.columns(2)
            
            # CSV export
            with export_col1:
                csv = df.to_csv(index=False).encode('utf-8')
                st.download_button("Download CSV", csv, "ec2_sizing.csv", "text/csv")
            
            # DOCX export
            def create_docx_report(df):
                doc = Document()
                
                # Title
                title = doc.add_paragraph()
                title_run = title.add_run(f"AWS EC2 SQL Server Sizing Report - {region}")
                title_run.font.size = Pt(18)
                title_run.font.bold = True
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Metadata
                doc.add_paragraph(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Input Parameters:")
                doc.add_paragraph(f"  CPU Cores: {inputs['on_prem_cores']}")
                doc.add_paragraph(f"  Peak CPU: {inputs['peak_cpu_percent']}%")
                doc.add_paragraph(f"  RAM: {inputs['on_prem_ram_gb']} GB")
                doc.add_paragraph(f"  Peak RAM: {inputs['peak_ram_percent']}%")
                doc.add_paragraph(f"  Storage: {inputs['storage_current_gb']} GB")
                doc.add_paragraph(f"  Growth Rate: {inputs['storage_growth_rate']*100:.1f}%")
                doc.add_paragraph(f"  Projection Years: {inputs['years']}")
                doc.add_paragraph(f"  Region: {inputs['region']}")
                
                # Table
                doc.add_heading("Recommendations", level=1)
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                
                # Header
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = col
                
                # Data rows with currency formatting
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, col in enumerate(df.columns):
                        value = row[col]
                        if col in ["instance_cost", "ebs_cost", "total_cost"]:
                            row_cells[i].text = f"${value:,.2f}"
                        else:
                            row_cells[i].text = str(value)
                
                # Recommendations
                doc.add_heading("Implementation Notes", level=1)
                notes = [
                    "PROD environments should use Multi-AZ deployments for high availability",
                    "Use gp3 EBS volumes for cost-effective general storage",
                    "Use io2 EBS volumes for high-performance needs (>16K IOPS or >1GB/s throughput)",
                    "Enable EBS encryption at rest for all environments",
                    "Implement regular snapshot backups with retention policies",
                    "Monitor performance metrics with Amazon CloudWatch",
                    "Consider Reserved Instances for PROD for cost savings"
                ]
                
                for note in notes:
                    p = doc.add_paragraph(style='ListBullet')
                    p.add_run(note)
                
                return doc
            
            with export_col2:
                doc = create_docx_report(df)
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                st.download_button("Download DOCX Report", doc_io, 
                                  "ec2_sizing_report.docx", 
                                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            # Cost optimization note
            st.info(f"💲 **Cost Estimates**: Monthly costs for {region} region include EC2 instance (Windows) and EBS storage. " 
                    "Actual costs may vary based on usage patterns and discounts.")
            
            if inputs["prefer_amd"]:
                st.info("💡 **Cost Optimization Tip**: AMD-based instances (m6a, r6a, c6a) typically offer 10-20% better price/performance than comparable Intel instances.")

        except Exception as e:
            st.error(f"Error generating recommendations: {str(e)}")

# Footer
st.markdown("---")
st.markdown("""
**Enterprise Features:**
- Multi-region cost estimates (us-east-1, us-west-1, us-west-2)
- Environment-specific sizing (DEV, QA, SQA, PROD)
- AMD instance optimization for cost savings
- Storage growth projections
- I/O requirements calculation
- Professional reporting (CSV, DOCX)
- Pricing validation and caching
- AWS credential management
""")